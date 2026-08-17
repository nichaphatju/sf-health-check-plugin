#!/usr/bin/env python3
import argparse
import json
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAVY = "0B356B"
BLUE = "168AC5"
DARK = "2C2C2C"
MID = "666666"
LIGHT = "F3F5F7"
BORDER = "D5D9DD"
CRITICAL = "8F1D1D"
HIGH = "B33A3A"
MEDIUM = "C97B13"
LOW = "347A58"
INFO = "4E6A7A"
WHITE = "FFFFFF"

SEVERITY_COLORS = {
    "Critical": CRITICAL,
    "High": HIGH,
    "Medium": MEDIUM,
    "Low": LOW,
    "Informational": INFO,
}

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def borders(cell, color=BORDER, size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tbl = tcPr.find(qn("w:tcBorders"))
    if tbl is None:
        tbl = OxmlElement("w:tcBorders")
        tcPr.append(tbl)
    for edge in ("top","left","bottom","right"):
        el = tbl.find(qn("w:"+edge))
        if el is None:
            el = OxmlElement("w:"+edge)
            tbl.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_cell_text(cell, text, bold=False, color=DARK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("" if text is None else str(text))
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        shade(hdr.cells[i], NAVY)
        borders(hdr.cells[i])
        set_cell_text(hdr.cells[i], h, bold=True, color=WHITE)
        if widths:
            hdr.cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            borders(cells[i])
            set_cell_text(cells[i], v)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph("")
    return table

def replace_tokens(doc, mapping):
    for p in list(doc.paragraphs) + [p for s in doc.sections for p in s.header.paragraphs] + [p for s in doc.sections for p in s.footer.paragraphs]:
        if not p.text:
            continue
        text = p.text
        changed = False
        for k, v in mapping.items():
            if k in text:
                text = text.replace(k, v or "")
                changed = True
        if changed:
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)

def remove_marker(doc, marker):
    for p in doc.paragraphs:
        if marker in p.text:
            p._element.getparent().remove(p._element)
            return

def add_h1(doc, text):
    p = doc.add_paragraph(text, style="HC H1")
    return p

def add_h2(doc, text):
    return doc.add_paragraph(text, style="HC H2")

def add_h3(doc, text):
    return doc.add_paragraph(text, style="HC H3")

def add_body(doc, text):
    if text:
        return doc.add_paragraph(text, style="HC Body")

def add_bullets(doc, items):
    for item in items or []:
        p = doc.add_paragraph(style="HC Body")
        p.style = doc.styles["HC Body"]
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.add_run("•  " + str(item))

def add_label_block(doc, label, text):
    p = doc.add_paragraph(style="HC Body")
    r = p.add_run(label + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text or "—")

def add_finding(doc, finding, source_map):
    color = SEVERITY_COLORS.get(finding["severity"], INFO)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(6.1)
    right.width = Inches(1.0)
    for c in (left, right):
        borders(c, color, "7")
    shade(left, "F7F9FB")
    shade(right, color)
    set_cell_text(left, f'{finding["id"]} — {finding["title"]}', bold=True, color=NAVY, size=11)
    set_cell_text(right, finding["severity"].upper(), bold=True, color=WHITE, size=9)
    doc.add_paragraph("")

    meta = [
        ("Confidence", finding["confidence"]),
        ("Guidance Basis", finding["guidanceBasis"]),
        ("Affected", ", ".join(finding.get("affectedComponents") or [])),
        ("Well-Architected", ", ".join(finding.get("wellArchitectedLens") or []) or "—"),
    ]
    add_table(doc, ["Field", "Value"], meta, widths=[1.5, 5.5])

    add_label_block(doc, "Observation", finding.get("observation"))

    evidence_texts = []
    for e in finding.get("evidence", []):
        bits = [e.get("type")]
        if e.get("component"): bits.append(e["component"])
        if e.get("location"): bits.append(e["location"])
        prefix = " | ".join([b for b in bits if b])
        evidence_texts.append((prefix + ": " if prefix else "") + e.get("detail",""))
    add_label_block(doc, "Org Evidence", "\n".join("• " + x for x in evidence_texts) if evidence_texts else "—")

    guidance = finding.get("guidance") or {}
    gtext = guidance.get("summary") or ""
    refs = []
    for sid in guidance.get("sourceIds") or []:
        src = source_map.get(sid)
        if src:
            refs.append(f'{src.get("title")} — {src.get("url")}')
    if refs:
        gtext += ("\n\nSources:\n" if gtext else "Sources:\n") + "\n".join("• " + r for r in refs)
    add_label_block(doc, "Salesforce Guidance", gtext or "No verified official guidance recorded.")

    add_label_block(doc, "Gap / Alignment", finding.get("gap"))
    add_label_block(doc, "Risk / Impact", finding.get("risk"))
    add_label_block(doc, "Recommendation", finding.get("recommendation"))

    p = doc.add_paragraph(style="HC Body")
    r = p.add_run(f'Priority: {finding["priority"]}     Effort: {finding["effort"]}')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    doc.add_paragraph("")

def phase_for(priority):
    if priority == "Immediate":
        return "Phase 1 — Immediate Risk Reduction"
    if priority == "Near-term":
        return "Phase 2 — Stabilisation"
    return "Phase 3 — Optimisation & Modernisation"

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("report")
    ap.add_argument("--template", required=True)
    ap.add_argument("--output", required=True)
    args = ap.parse_args()

    data = json.loads(Path(args.report).read_text(encoding="utf-8"))
    doc = Document(args.template)

    a = data["assessment"]
    mapping = {
        "{{ORGANISATION}}": a.get("organisation",""),
        "{{ORG_ENVIRONMENT}}": " / ".join([x for x in [a.get("orgName"), a.get("orgType")] if x]),
        "{{EDITION}}": a.get("edition",""),
        "{{ORG_ID}}": a.get("orgId",""),
        "{{EVIDENCE_DATE}}": a.get("evidenceDate",""),
        "{{ASSESSMENT_DATE}}": a.get("assessmentDate",""),
        "{{PREPARED_BY}}": a.get("preparedBy",""),
        "{{VERSION}}": a.get("version",""),
        "{{CLASSIFICATION}}": a.get("classification","CONFIDENTIAL"),
    }
    replace_tokens(doc, mapping)
    remove_marker(doc, "{{REPORT_BODY_START}}")
    doc.add_page_break()

    add_h1(doc, "Contents")
    contents = [
        "1. Executive Summary",
        "2. Assessment Overview",
        "3. Org Health Snapshot",
        "4. Key Risks & Findings",
        "5. Detailed Assessment",
        "6. Remediation Roadmap",
        "7. Positive Findings",
    ]
    if data.get("componentHighlights"):
        contents.extend([
            "8. Selected Component Deep Dives",
            "9. Assessment Limitations",
        ])
    else:
        contents.append("8. Assessment Limitations")
    contents.extend([
        "Appendix A. Detailed Finding Register",
        "Appendix B. Org Inventory / Metrics",
        "Appendix C. Evidence Sources",
        "Appendix D. Official Salesforce Guidance References",
    ])
    for item in contents:
        p = doc.add_paragraph(item, style="HC Body")
        p.paragraph_format.space_after = Pt(5)
    doc.add_page_break()

    findings = data.get("findings", [])
    sev = Counter(f["severity"] for f in findings)
    by_cat = defaultdict(list)
    for f in findings:
        by_cat[f["category"]].append(f)
    source_map = {s["id"]: s for s in data.get("salesforceSources", [])}

    add_h1(doc, "1. Executive Summary")
    summary = data["executiveSummary"]
    p = doc.add_paragraph(style="HC H2")
    p.add_run("Overall Assessment: ")
    rr = p.add_run(summary["overallAssessment"])
    rr.bold = True
    add_body(doc, summary["summary"])
    add_h2(doc, "Key Themes")
    add_bullets(doc, summary.get("keyThemes", []))
    if summary.get("recommendedNextActions"):
        add_h2(doc, "Recommended Next Actions")
        add_bullets(doc, summary["recommendedNextActions"])

    add_h1(doc, "2. Assessment Overview")
    add_h2(doc, "Scope")
    add_bullets(doc, a.get("scope", []))
    add_h2(doc, "Methodology")
    add_body(doc, "The assessment uses an evidence-first process. Org observations are identified from the reviewed Salesforce evidence, then material best-practice conclusions are validated against current official Salesforce guidance before being promoted to findings.")
    add_h2(doc, "Evidence Reviewed")
    evrows = []
    for e in data.get("evidenceSources", []):
        evrows.append([e.get("id",""), e.get("name",""), e.get("type",""), e.get("date","")])
    if evrows:
        add_table(doc, ["ID","Evidence","Type","Date"], evrows, widths=[0.7,3.5,1.7,1.1])
    add_h2(doc, "Assessment Limitations")
    add_bullets(doc, data.get("limitations", []))

    add_h1(doc, "3. Org Health Snapshot")
    covrows = []
    for c in data.get("coverage", []):
        cat_findings = by_cat.get(c["area"], [])
        counts = Counter(f["severity"] for f in cat_findings)
        covrows.append([
            c["area"],
            c["status"],
            counts.get("Critical",0),
            counts.get("High",0),
            counts.get("Medium",0),
            counts.get("Low",0),
            c["summary"]
        ])
    if covrows:
        add_table(doc, ["Area","Assessment","Crit","High","Med","Low","Summary"], covrows, widths=[1.2,1.25,0.42,0.42,0.42,0.42,3.0])

    add_h1(doc, "4. Key Risks & Findings")
    key = [f for f in findings if f["severity"] in {"Critical","High"}]
    if not key:
        key = [f for f in findings if f["severity"] == "Medium"][:8]
    keyrows = [[f["id"], f["title"], f["severity"], f["confidence"], f["category"], f["priority"]] for f in key]
    if keyrows:
        add_table(doc, ["ID","Finding","Severity","Confidence","Area","Priority"], keyrows, widths=[0.8,2.65,0.75,1.0,1.0,0.9])
    else:
        add_body(doc, "No material Critical, High or selected systemic Medium findings were recorded.")

    add_h1(doc, "5. Detailed Assessment")
    for category in sorted(by_cat.keys()):
        add_h2(doc, category)
        for f in by_cat[category]:
            add_finding(doc, f, source_map)

    add_h1(doc, "6. Remediation Roadmap")
    phases = defaultdict(list)
    for f in findings:
        if f["severity"] != "Informational":
            phases[phase_for(f["priority"])].append(f)
    for phase in ["Phase 1 — Immediate Risk Reduction", "Phase 2 — Stabilisation", "Phase 3 — Optimisation & Modernisation"]:
        add_h2(doc, phase)
        rows = [[f["id"], f["recommendation"], f["priority"], f["effort"]] for f in phases.get(phase, [])]
        if rows:
            add_table(doc, ["Finding","Recommended Action","Priority","Effort"], rows, widths=[0.9,4.9,1.0,0.6])
        else:
            add_body(doc, "No findings assigned to this phase.")

    add_h1(doc, "7. Positive Findings")
    positives = data.get("positiveFindings", [])
    if positives:
        for pf in positives:
            add_h2(doc, pf["title"])
            add_body(doc, pf["summary"])
    else:
        add_body(doc, "No material positive findings were recorded separately.")

    highlights = data.get("componentHighlights", [])
    if highlights:
        add_h1(doc, "8. Selected Component Deep Dives")
        for h in highlights:
            add_h2(doc, f'{h["component"]} — {h["type"]}')
            add_body(doc, h["reason"])
            if h.get("findingIds"):
                add_body(doc, "Related findings: " + ", ".join(h["findingIds"]))
            if h.get("diagramPath"):
                pth = Path(h["diagramPath"])
                if pth.exists():
                    doc.add_picture(str(pth), width=Inches(6.3))
                    cap = doc.add_paragraph(style="HC Small")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.add_run(f'Figure — {h["component"]}')

    add_h1(doc, "9. Assessment Limitations" if highlights else "8. Assessment Limitations")
    add_bullets(doc, data.get("limitations", []))
    add_body(doc, "This assessment reflects the Salesforce evidence available as of the stated evidence date. Static analysis and metadata review cannot establish the absence of defects, security vulnerabilities, runtime issues, or unused functionality outside the evidence reviewed.")

    add_h1(doc, "Appendix A. Detailed Finding Register")
    regrows = [[f["id"], f["severity"], f["confidence"], f["category"], f["title"], f["priority"]] for f in findings]
    if regrows:
        add_table(doc, ["ID","Severity","Confidence","Category","Finding","Priority"], regrows, widths=[0.85,0.75,1.05,1.1,2.7,0.85])

    add_h1(doc, "Appendix B. Org Inventory / Metrics")
    metrics = data.get("metrics", {})
    if metrics:
        add_table(doc, ["Metric","Value"], [[k,v] for k,v in metrics.items()], widths=[3.2,3.8])
    else:
        add_body(doc, "No additional measured metrics were supplied.")

    add_h1(doc, "Appendix C. Evidence Sources")
    if data.get("evidenceSources"):
        add_table(doc, ["ID","Name","Type","Reference","Date"], [
            [e.get("id",""), e.get("name",""), e.get("type",""), e.get("pathOrReference",""), e.get("date","")]
            for e in data["evidenceSources"]
        ], widths=[0.6,2.3,1.25,2.2,0.8])

    add_h1(doc, "Appendix D. Official Salesforce Guidance References")
    if data.get("salesforceSources"):
        for s in data["salesforceSources"]:
            add_h2(doc, f'{s["id"]} — {s["title"]}')
            add_body(doc, s["url"])
            if s.get("findingIds"):
                add_body(doc, "Supports: " + ", ".join(s["findingIds"]))
    else:
        add_body(doc, "No official Salesforce sources were recorded.")

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)

if __name__ == "__main__":
    main()
